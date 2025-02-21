import streamlit as st
from docx import Document
import io
import os
from datetime import datetime

# Function to generate the resume
def generate_resume(name, phone, email, github, linkedin, dob, gender, location, languages, experience, internships, projects, skills, education, job_title, objective, summary, achievements, certifications, profile_photo):
    doc = Document()
    
    doc.add_heading(f'{job_title} Resume', 0)
    
    # Profile Section
    doc.add_heading('Profile:', level=1)
    if profile_photo:
        doc.add_paragraph("Profile Photo Attached")
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Date of Birth: {dob}")
    doc.add_paragraph(f"Gender: {gender}")
    doc.add_paragraph(f"Location: {location}")
    doc.add_paragraph(f"Phone: {phone}")
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"GitHub: {github}")
    doc.add_paragraph(f"LinkedIn: {linkedin}")
    doc.add_paragraph(f"Languages: {languages}")
    
    doc.add_heading('Professional Summary:', level=1)
    doc.add_paragraph(summary)
    
    doc.add_heading('Job Objective:', level=1)
    doc.add_paragraph(objective)
    
    doc.add_heading('Experience:', level=1)
    for exp in experience:
        doc.add_paragraph(f"{exp['position']} at {exp['company']} ({exp['start_date']} - {exp['end_date']})")
        doc.add_paragraph(exp['responsibilities'])
    
    doc.add_heading('Internships:', level=1)
    for intern in internships:
        doc.add_paragraph(f"{intern['position']} at {intern['company']} ({intern['start_date']} - {intern['end_date']})")
        doc.add_paragraph(intern['responsibilities'])
    
    doc.add_heading('Projects:', level=1)
    for proj in projects:
        doc.add_paragraph(f"{proj['title']} - {proj['description']}")
    
    doc.add_heading('Education:', level=1)
    doc.add_paragraph(education)
    
    doc.add_heading('Skills:', level=1)
    if skills:
        for skill in skills:
            doc.add_paragraph(skill)
    
    doc.add_heading('Achievements:', level=1)
    doc.add_paragraph(achievements)
    
    doc.add_heading('Certifications:', level=1)
    doc.add_paragraph(certifications)
    
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

st.set_page_config(page_title="AI Resume Generator", layout="wide")

st.sidebar.subheader("Login")
username = st.sidebar.text_input("Username")
password = st.sidebar.text_input("Password", type="password")
if username and password:
    st.sidebar.success(f"Logged in as {username}")
else:
    st.sidebar.warning("Please log in to continue.")
    st.stop()

st.title("📄 AI Resume Generator")
profile_photo = st.file_uploader("Upload Profile Photo (optional)", type=["png", "jpg", "jpeg"])
name = st.text_input("Full Name")
dob = st.date_input("Date of Birth", min_value=datetime(1950, 1, 1))  # FIXED ERROR HERE
gender = st.selectbox("Gender", ["Male", "Female", "Other"])
location = st.text_input("Location")
phone = st.text_input("Phone Number", max_chars=15)
email = st.text_input("Email Address")
github = st.text_input("GitHub Profile Link")
linkedin = st.text_input("LinkedIn Profile Link")
languages = st.text_input("Languages Known (comma-separated)")
job_title = st.text_input("Target Job Title")
summary = st.text_area("Professional Summary")
objective = st.text_area("Job Objective")
achievements = st.text_area("Achievements")
education = st.text_area("Education Details")
certifications = st.text_area("Certifications")

experience = []
with st.expander("➕ Add Work Experience"):
    num_experience = st.number_input("Number of Experiences", min_value=0, step=1)
    for i in range(num_experience):
        position = st.text_input(f"Position {i+1}", key=f"pos_{i}")
        company = st.text_input(f"Company {i+1}", key=f"comp_{i}")
        start_date = st.date_input(f"Start Date {i+1}", key=f"start_{i}")
        end_date = st.selectbox(f"End Date {i+1}", ["Presently Working"] + [str(y) for y in range(1980, 2025)], key=f"end_{i}")
        responsibilities = st.text_area(f"Responsibilities {i+1}", key=f"resp_{i}")
        experience.append({"position": position, "company": company, "start_date": str(start_date), "end_date": end_date, "responsibilities": responsibilities})

internships = []
with st.expander("➕ Add Internships"):
    num_internships = st.number_input("Number of Internships", min_value=0, step=1)
    for i in range(num_internships):
        position = st.text_input(f"Internship {i+1}", key=f"intern_{i}")
        company = st.text_input(f"Company {i+1}", key=f"intern_comp_{i}")
        start_date = st.date_input(f"Start Date {i+1}", key=f"intern_start_{i}")
        end_date = st.selectbox(f"End Date {i+1}", ["Presently Working"] + [str(y) for y in range(1980, 2025)], key=f"intern_end_{i}")
        responsibilities = st.text_area(f"Responsibilities {i+1}", key=f"intern_resp_{i}")
        internships.append({"position": position, "company": company, "start_date": str(start_date), "end_date": end_date, "responsibilities": responsibilities})

projects = []
with st.expander("➕ Add Projects"):
    num_projects = st.number_input("Number of Projects", min_value=0, step=1)
    for i in range(num_projects):
        title = st.text_input(f"Project {i+1}", key=f"proj_{i}")
        description = st.text_area(f"Project Description {i+1}", key=f"proj_desc_{i}")
        projects.append({"title": title, "description": description})

skills = st.text_area("Enter Skills (comma-separated)")
skills = [skill.strip() for skill in skills.split(",")] if skills else []

if st.button("📥 Generate Resume"):
    resume_file = generate_resume(name, phone, email, github, linkedin, dob, gender, location, languages, experience, internships, projects, skills, education, job_title, objective, summary, achievements, certifications, profile_photo)
    st.download_button("⬇️ Download Resume", resume_file, file_name="resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.sidebar.markdown("---")
mode = st.sidebar.radio("Theme Mode", ["🌞 Light", "🌙 Dark"], key="theme_mode")
